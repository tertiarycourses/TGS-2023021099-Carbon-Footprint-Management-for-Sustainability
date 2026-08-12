#!/usr/bin/env python3
"""Generate four assessment instruments aligned to the approved legacy assessment plan."""
import os, sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__))
REPO=os.path.dirname(HERE)
sys.path.insert(0,HERE)
import course_data as C
import prodoc

OUT=os.path.join(REPO,"assessment")
LOGO=os.path.join(REPO,"courseware","assets","tertiary-infotech-logo.png")
os.makedirs(OUT,exist_ok=True)

WA=[
 ("K1",8,"Explain how ISO 14001, ISO 14064-1 and the GHG Protocol Corporate Standard contribute different but complementary controls to an organisational carbon-footprint programme."),
 ("K2",8,"Describe the Plan-Do-Check-Act cycle and give one carbon-management action for each stage."),
 ("K3",10,"A company’s screening inventory is: purchased materials 4,200 tCO₂e, electricity 2,800 tCO₂e, freight 900 tCO₂e, fuel 1,200 tCO₂e and waste 500 tCO₂e. Identify the two largest hotspots, calculate their combined percentage of the total, and state one further factor to consider before prioritising them."),
 ("K4",10,"State the purpose of a business-as-usual emissions scenario. Then describe how activity growth, efficiency improvement and emission-factor change should be represented in a projection."),
 ("K5",10,"Recommend two technically different carbon-reduction measures for a manufacturing facility and explain one feasibility test and one operational guardrail for each."),
 ("K6",12,"List and explain six essential components of an organisational carbon-footprint reduction plan."),
 ("K7",12,"A retrofit costs S$150,000 and is expected to save S$42,000 and 85 tCO₂e each year for eight years. Calculate the simple payback and the undiscounted lifetime marginal abatement cost. Interpret both results."),
 ("K8",10,"Explain how a multi-criteria decision analysis can combine carbon impact, financial value, feasibility, strategic fit and data confidence. State one limitation of using only the final weighted score."),
 ("K9",10,"Draft the fields that should appear in a carbon-reduction action register and explain how a stage-gate review supports delivery and continual improvement."),
]

CASE="""EcoFab Components Pte Ltd manufactures precision parts at one Singapore facility. Its 2025 screening footprint is shown below. Production is expected to grow by 4% each year to 2030.

• Purchased aluminium: 4,500 tCO₂e (Scope 3; low data quality; medium controllability)
• Purchased electricity: 3,000 tCO₂e (Scope 2; high data quality; high controllability)
• Natural gas: 1,350 tCO₂e (Scope 1; high data quality; medium controllability)
• Freight: 850 tCO₂e (Scope 3; medium data quality; medium controllability)
• Waste: 300 tCO₂e (Scope 3; medium data quality; high controllability)

The board has approved an aspiration to reduce absolute emissions by 30% by 2030 from the 2025 baseline. Management has shortlisted the following initiatives; all figures are illustrative screening estimates.

1. Energy controls — S$180,000 capex; 450 tCO₂e/year; high feasibility; can start in 2026.
2. Rooftop solar — S$650,000 capex; 480 tCO₂e/year; medium feasibility; requires a roof study.
3. Heat recovery — S$420,000 capex; 330 tCO₂e/year; medium feasibility; requires a shutdown window.
4. Low-carbon aluminium — S$300,000 annual premium; 1,500 tCO₂e/year; low data confidence; requires supplier product-footprint evidence.
5. Freight optimisation — S$90,000 capex; 240 tCO₂e/year; high feasibility; requires carrier route data.
6. Waste-yield programme — S$120,000 capex; 180 tCO₂e/year; high feasibility; can start in 2026.

The first-wave capital budget is S$1,000,000. The board requires a plan that is quantified, implementable and transparent about assumptions."""

CS=[
 ("A1",30,"Analyse EcoFab’s carbon-emission hotspots. Show the total footprint, percentage contribution of each source, the two leading hotspots, and the data-quality improvements needed before a final investment decision."),
 ("A2",35,"Prioritise the six reduction initiatives. Use carbon impact, cost, feasibility, dependencies and data confidence; select a first-wave portfolio within the S$1,000,000 capital budget and justify funded, pilot and deferred choices."),
 ("A3",35,"Develop a 2026–2030 carbon-footprint reduction plan for EcoFab. Include the precise target, annual milestones or trajectory, initiative sequence, accountable roles, resources, KPIs, measurement and verification, management-review cycle, risks and communication controls."),
]

WA_KEY={
 "K1":["ISO 14001: EMS governance, aspects, objectives, controls, audit and PDCA.","ISO 14064-1: organisation-level GHG quantification and reporting requirements.","GHG Protocol: corporate boundary, scopes, principles and inventory method.","Explains complementarity rather than treating the standards as interchangeable."],
 "K2":["Plan: boundary, baseline, risks, target and actions.","Do: implement controls, projects and competence.","Check: monitor KPIs, audit and explain variance.","Act: corrective action, management review and plan revision."],
 "K3":["Total = 9,600 tCO₂e.","Materials = 43.75%; electricity = 29.17%; combined = 72.92%.","Further judgement may include controllability, risk, trend, data quality or strategic fit."],
 "K4":["BAU is the no-additional-action reference for comparing interventions.","Future activity reflects volume/growth; efficiency changes emissions per activity; factor change reflects energy/supplier system.","Assumptions, period, units and uncertainty cases are stated."],
 "K5":["Any two distinct credible measures, such as controls/efficiency, electrification, renewable energy, heat recovery, circularity or supplier substitution.","Feasibility covers technical, operational, financial and evidence dimensions.","Guardrails protect safety, quality, uptime, yield or service."],
 "K6":["Any six well-explained: boundary/baseline, hotspots, target, initiatives, roadmap, governance, resources, KPIs/M&V, risk, communication, review."],
 "K7":["Payback = 150,000 ÷ 42,000 = 3.57 years.","Lifetime net cost = 150,000 − (8 × 42,000) = −S$186,000.","Lifetime reduction = 680 tCO₂e; MAC = −186,000 ÷ 680 = −S$273.53/tCO₂e.","Interprets negative MAC as screened cost-saving; notes discounting/assumption limits."],
 "K8":["Agree criteria, weights and 1–5 evidence anchors; weighted sum supports comparison.","Sensitivity tests show whether ranking depends on weights.","Score does not replace judgement about dependencies, thresholds, uncertainty or governance."],
 "K9":["Register fields include ID, action, owner, dates, budget, benefit, dependency, KPI, baseline, M&V, status, issue and decision.","Stage gate compares evidence with pre-agreed scale/modify/stop criteria and records learning."],
}

CS_KEY={
 "A1":["Total footprint = 10,000 tCO₂e.","Percentages: aluminium 45%; electricity 30%; natural gas 13.5%; freight 8.5%; waste 3%.","Aluminium and electricity = 75% combined.","Prioritises supplier-specific aluminium PCF/boundary evidence and validates material/factor data; explains controllability and risk."],
 "A2":["Uses a transparent comparison rather than ranking only by tonnes.","A defensible first wave can fund energy controls, rooftop solar, freight optimisation and waste yield: S$1,040,000 exceeds cap, so one must be staged; for example fund controls + solar + freight = S$920,000, pilot waste yield from operating budget or defer subject to approval.","Low-carbon aluminium should be piloted/contracted subject to product-footprint evidence; heat recovery sequenced to shutdown window.","Accept alternative portfolios within cap where assumptions, dependencies and trade-offs are coherent."],
 "A3":["Target: reduce absolute organisational footprint 30% from 10,000 to 7,000 tCO₂e by 2030, with boundary and recalculation policy.","BAU growth is considered; selected initiatives quantify a credible pathway and do not double count.","Plan names executive accountability, carbon/GHG manager and process owners.","Includes initiative register, resources, annual/quarterly milestones, outcome/driver/delivery/data-quality KPIs, M&V, risks, management review and transparent claims."]
}

def setup_doc(kind,title):
    d=Document(); sec=d.sections[0]
    sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    d.styles["Normal"].font.name="Arial"; d.styles["Normal"].font.size=Pt(10.5)
    prodoc.style_headings(d); prodoc.add_page_numbers(d); prodoc.enable_update_fields(d)
    prodoc.add_cover_page(d,kind,title,C.VERSION,org_logo=LOGO,course_code=C.COURSE_CODE)
    return d

def instructions(d,instrument,duration,total):
    d.add_heading("Candidate Instructions and Grading",level=1)
    for label,value in [("Instrument",instrument),("Duration",duration),("Conditions","Open book; individual work"),("Total marks",str(total)),("Outcome","Competent (C) or Not Yet Competent (NYC), according to the approved assessment decision rules")]:
        p=d.add_paragraph(); p.add_run(label+": ").bold=True; p.add_run(value)
    d.add_heading("Instructions",level=2)
    for x in ["Answer every question.","Show calculations, units and assumptions where applicable.","Use only approved course materials and work independently.","Do not photograph, record or discuss the assessment.","If information is missing, state a reasonable assumption and explain its effect.","Submit as instructed by the assessor before time is called."]:
        d.add_paragraph(x,style="List Bullet")
    d.add_heading("Grading approach",level=2)
    d.add_paragraph("Marks are awarded for technically correct content, relevant application, transparent calculation and a defensible recommendation. The assessor applies the approved competency decision rules and records feedback.")
    d.add_page_break()

def identity(d):
    t=d.add_table(rows=2,cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.cell(0,0).text="Candidate name:"; t.cell(0,1).text="NRIC / ID (last 4):"
    t.cell(1,0).text="Assessment date:"; t.cell(1,1).text="Assessor:"
    d.add_paragraph()

def build_wa(answer=False):
    kind="ANSWER KEY — WRITTEN ASSESSMENT" if answer else "WRITTEN ASSESSMENT"
    d=setup_doc(kind,C.COURSE_TITLE)
    instructions(d,"Short-Answer Questions (9 questions, K1–K9)","1 hour",90)
    identity(d); d.add_heading("Short-Answer Questions" if not answer else "Model Answers and Marking Guidance",level=1)
    for i,(k,marks,q) in enumerate(WA,1):
        d.add_heading(f"Question {i} — {k} ({marks} marks)",level=2); d.add_paragraph(q)
        if answer:
            for item in WA_KEY[k]: d.add_paragraph(item,style="List Bullet")
            d.add_paragraph("Award proportionately for technically correct equivalent responses. Do not require verbatim wording.")
        else:
            for _ in range(4 if marks<=8 else 6): d.add_paragraph("________________________________________________________________________________")
    name=("Answers to WA (SAQ)" if answer else "WA (SAQ)")+f" - {C.COURSE_TITLE} - {C.VERSION}.docx"
    path=os.path.join(OUT,name); d.save(path); return path

def build_cs(answer=False):
    kind="ANSWER KEY — CASE STUDY" if answer else "CASE STUDY ASSESSMENT"
    d=setup_doc(kind,C.COURSE_TITLE)
    instructions(d,"Written Assessment — Case Study (3 questions, A1–A3)","1 hour",100)
    identity(d); d.add_heading("Case Study: EcoFab Components Pte Ltd",level=1); d.add_paragraph(CASE)
    d.add_page_break(); d.add_heading("Case Study Questions" if not answer else "Model Answers and Marking Guidance",level=1)
    for i,(a,marks,q) in enumerate(CS,1):
        if i>1 and not answer: d.add_page_break()
        d.add_heading(f"Question {i} — {a} ({marks} marks)",level=2); d.add_paragraph(q)
        if answer:
            for item in CS_KEY[a]: d.add_paragraph(item,style="List Bullet")
            d.add_paragraph("Accept an alternative recommendation where calculations are correct and the evidence, constraints, dependencies and limitations are coherently defended.")
        else:
            for _ in range(16): d.add_paragraph("________________________________________________________________________________")
    name=("Answers to CS Assessment" if answer else "CS Assessment")+f" - {C.COURSE_TITLE} - {C.VERSION}.docx"
    path=os.path.join(OUT,name); d.save(path); return path

if __name__=="__main__":
    for p in [build_wa(False),build_wa(True),build_cs(False),build_cs(True)]: print("Saved",p)
