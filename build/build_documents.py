#!/usr/bin/env python3
"""Generate the Learner Guide and Lesson Plan from the same source model."""
import json, os, sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE=os.path.dirname(os.path.abspath(__file__))
REPO=os.path.dirname(HERE)
sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1, SLIDES1
from data_domain2 import DOMAIN2, SLIDES2
from data_domain3 import DOMAIN3, SLIDES3
import prodoc

OUT=os.path.join(REPO,"courseware")
LOGO=os.path.join(OUT,"assets","tertiary-infotech-logo.png")

def base_doc():
    d=Document(); sec=d.sections[0]
    sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.76); sec.right_margin=Inches(.76)
    styles=d.styles
    styles["Normal"].font.name="Arial"; styles["Normal"].font.size=Pt(10.5)
    styles["Normal"].paragraph_format.space_after=Pt(5)
    prodoc.style_headings(d); prodoc.add_page_numbers(d); prodoc.enable_update_fields(d)
    return d

def heading(d,text,level=1):
    d.add_heading(text,level=level)

def bullets(d,items):
    for item in items: d.add_paragraph(str(item),style="List Bullet")

def add_source_note(d,text):
    p=d.add_paragraph(); r=p.add_run("Evidence note: "); r.bold=True
    p.add_run(text)

def add_activity(d,a):
    heading(d,f"Activity {a['id']}: {a['title']}",2)
    t=d.add_table(rows=0,cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for k,v in [("Real-use-case lens",a["case"]),("Source lens",a["source"]),("Objective",a["objective"]),("Deliverable",a["deliverable"])]:
        c=t.add_row().cells; c[0].text=k; c[1].text=v
    heading(d,"Scenario",3); d.add_paragraph(a["scenario"])
    add_source_note(d,"Company-linked facts are used only as a public case lens. Unless explicitly identified as reported, all activity figures are illustrative teaching data.")
    heading(d,"Detailed procedure",3)
    for i,step in enumerate(a["steps"],1): d.add_paragraph(f"{i}. {step}")
    heading(d,"Scenario questions",3)
    for i,q in enumerate(a["questions"],1): d.add_paragraph(f"{i}. {q}")
    heading(d,"Evidence and reflection",3)
    bullets(d,["State the boundary, period, unit, factor and source.","Show calculations and retain intermediate values.","Separate evidence, assumption and judgement.","Record uncertainty and the evidence that would change the decision.","Submit the specified deliverable and a concise management recommendation."])

def build_lg():
    d=base_doc(); prodoc.add_cover_page(d,"Learner Guide",C.COURSE_TITLE,C.VERSION,org_logo=LOGO,course_code=C.COURSE_CODE)
    prodoc.add_version_control(d,[(C.VERSION,C.VERSION_DATE,"Rebuilt from legacy v6; expanded standards, Singapore context, real-use-case activities and assessment alignment",C.TRAINER)])
    prodoc.add_toc(d)
    heading(d,"How to Use This Learner Guide",1)
    d.add_paragraph("This guide contains the detailed concepts, activity procedures, scenarios and questions. The slide deck is deliberately visual and concise; use this guide for procedural detail, calculations and evidence expectations.")
    heading(d,"Course Information",2)
    for label,value in [("Course",C.COURSE_TITLE),("Course code",C.COURSE_CODE),("TSC",f"{C.TSC_TITLE} — {C.TSC_CODE}"),("Duration",C.COURSE_HOURS),("Trainer",C.TRAINER)]:
        p=d.add_paragraph(); p.add_run(label+": ").bold=True; p.add_run(value)
    heading(d,"Learning Outcomes",2); bullets(d,C.LEARNING_OUTCOMES)
    heading(d,"Assessment",2); bullets(d,[C.ASSESSMENT["wa"],C.ASSESSMENT["cs"],"Both instruments are open book and completed independently."])
    for t,specs,activities in zip(C.TOPICS,[SLIDES1,SLIDES2,SLIDES3],[DOMAIN1,DOMAIN2,DOMAIN3]):
        d.add_page_break(); heading(d,f"Topic {t['code']}: {t['title']}",1)
        d.add_paragraph(t["subtitle"]); p=d.add_paragraph(); p.add_run("Competency alignment: ").bold=True; p.add_run(f"{t['ks']} · {t['a']}")
        for spec in specs:
            heading(d,spec["title"],2)
            typ=spec["type"]
            if typ=="grid":
                for item in spec["items"]:
                    p=d.add_paragraph();
                    if isinstance(item,tuple): p.add_run(item[0]+": ").bold=True; p.add_run(item[1])
                    else: p.add_run(str(item))
            elif typ=="process":
                for i,step in enumerate(spec["steps"],1):
                    if isinstance(step,tuple): step=" — ".join(step)
                    d.add_paragraph(f"{i}. {step}")
            elif typ=="compare":
                table=d.add_table(rows=1,cols=2); table.style="Table Grid"
                table.rows[0].cells[0].text=spec["left_title"]; table.rows[0].cells[1].text=spec["right_title"]
                for i in range(max(len(spec["left"]),len(spec["right"]))):
                    c=table.add_row().cells; c[0].text=spec["left"][i] if i<len(spec["left"]) else ""; c[1].text=spec["right"][i] if i<len(spec["right"]) else ""
            elif typ=="decision":
                d.add_paragraph(spec["question"])
                p=d.add_paragraph(); p.add_run("If yes / fit: ").bold=True; p.add_run(spec["yes"])
                p=d.add_paragraph(); p.add_run("If no / not yet: ").bold=True; p.add_run(spec["no"])
                d.add_paragraph(spec.get("note",""))
            elif typ=="chart":
                table=d.add_table(rows=1,cols=2); table.style="Table Grid"; table.rows[0].cells[0].text="Category"; table.rows[0].cells[1].text="Illustrative value"
                for cat,val in zip(spec["categories"],spec["values"]): c=table.add_row().cells; c[0].text=str(cat); c[1].text=str(val)
                d.add_paragraph(spec.get("note",""))
            elif typ in ("formula","worked"):
                d.add_paragraph(spec["intro"]); p=d.add_paragraph(); r=p.add_run(spec["formula"]); r.bold=True
                bullets(d,spec["notes"])
            elif typ=="statement":
                p=d.add_paragraph(); r=p.add_run(spec["title"]); r.bold=True; r.font.size=Pt(14); d.add_paragraph(spec.get("sub",""))
        heading(d,f"Topic {t['code']} Activities",1)
        for a in activities: add_activity(d,a)
    d.add_page_break(); heading(d,"Assessment Preparation",1)
    bullets(d,["Review K1–K9 definitions and be able to apply each to a short scenario.","Practise showing units and calculations clearly.","For the case study, structure answers as evidence → analysis → recommendation → limitation.","Do not invent missing company facts; state a reasonable assumption and its effect."])
    heading(d,"Source Register",1)
    for name,url in C.SOURCES:
        p=d.add_paragraph(); p.add_run(name+": ").bold=True; p.add_run(url)
    heading(d,"Important Use Note",2)
    d.add_paragraph("Regulatory rates, thresholds, factors and reporting requirements can change. Verify current primary-source requirements before making external claims, filing reports or committing capital.")
    path=os.path.join(OUT,f"LG-{C.SHORT_TITLE}-{C.VERSION}.docx"); d.save(path); return path

def add_lp_row(table,cells):
    row=table.add_row().cells
    for i,val in enumerate(cells): row[i].text=str(val); row[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP

def build_lp():
    with open(os.path.join(HERE,"slide_map.json"),encoding="utf-8") as f: sm=json.load(f)
    d=base_doc(); prodoc.add_cover_page(d,"LESSON PLAN",C.COURSE_TITLE,C.VERSION,org_logo=LOGO,course_code=C.COURSE_CODE)
    prodoc.add_version_control(d,[(C.VERSION,C.VERSION_DATE,"Rebuilt and mapped to the v7 133-slide deck",C.TRAINER)])
    prodoc.add_toc(d)
    heading(d,"Course Delivery Overview",1)
    d.add_paragraph("Two training days of 8 hours each. Assessment is scheduled separately for 2 hours. Break durations are not counted as training time.")
    heading(d,"Facilitation Principles",2); bullets(d,["Use the slides for visual explanation and the Learner Guide for procedures.","Keep company-linked figures clearly labelled as illustrative unless a source explicitly reports them.","Require evidence, units, assumptions and limitations in every activity output.","Use peer review before trainer synthesis."])
    cols=["Time","Duration","Learning segment","Slides","Method and trainer actions","Learner evidence"]
    schedules=[
      (1,[
       ("09:00–09:45","0:45","Administration, introductions and outcomes","1–17","Briefing, introductions and carbon-management decision cycle","Expectations and prior-experience map"),
       ("09:45–10:45","1:00","K1 standards and inventory framing",f"{sm['topics']['1'][0]+1}–{sm['topics']['1'][0]+9}","Visual explanation, boundary checks and worked calculation","Boundary and principle responses"),
       ("10:45–11:00","—","Break","—","Not counted as training time","—"),
       ("11:00–12:30","1:30","K3 scopes, data and hotspot diagnosis",f"{sm['topics']['1'][0]+10}–{sm['topics']['1'][0]+18}","Scope classification, Pareto reading and calculator-fit decision","Scope map and hotspot interpretation"),
       ("12:30–13:30","—","Lunch","—","Not counted as training time","—"),
       ("13:30–14:30","1:00","K7 economics and investment tools",f"{sm['topics']['1'][0]+19}–{sm['topics']['1'][0]+27}","Payback, NPV, MAC and sensitivity examples","Calculation trail and economic caveats"),
       ("14:30–17:15","2:45","Activities 1–3",f"{sm['activities']['1']}–{sm['activities']['3']}","Team analysis, tool demonstration, calculation and presentations","Boundary memo, engagement experiment and hotspot recommendation"),
       ("17:15–18:15","1:00","Topic synthesis and formative check",f"{sm['activities']['1']-3}–{sm['activities']['1']-1}","Peer critique, trainer synthesis and individual exit response","Corrected decision trail and exit ticket")]),
      (2,[
       ("09:00–09:15","0:15","Recap and outcomes","—","Retrieval questions and agenda","Recall check"),
       ("09:15–10:30","1:15","K4 projections and scenarios",f"{sm['topics']['2'][0]+1}–{sm['topics']['2'][0]+9}","Scenario modelling and fleet worked example","Scenario assumptions register"),
       ("10:30–10:45","—","Break","—","Not counted as training time","—"),
       ("10:45–12:00","1:15","K5 technologies and feasibility",f"{sm['topics']['2'][0]+10}–{sm['topics']['2'][0]+18}","Technology cards, feasibility gate and pilot design","Feasibility scorecard"),
       ("12:00–13:00","1:00","K8 portfolio prioritisation",f"{sm['topics']['2'][0]+19}–{sm['topics']['2'][0]+27}","MCDA, MACC, dependencies and bias checks","Portfolio rationale"),
       ("13:00–14:00","—","Lunch","—","Not counted as training time","—"),
       ("14:00–15:30","1:30","Activities 4–6",f"{sm['activities']['4']}–{sm['activities']['6']}","Scenario model, pilot screen and portfolio workshop","Model, charter and decision record"),
       ("15:30–16:30","1:00","K2 EMS and governance",f"{sm['topics']['3'][0]+1}–{sm['topics']['3'][0]+9}","PDCA, controls and variance decision","Governance and control map"),
       ("16:30–16:45","—","Break","—","Not counted as training time","—"),
       ("16:45–17:45","1:00","K6 and K9 plan architecture",f"{sm['topics']['3'][0]+10}–{sm['topics']['3'][0]+27}","Target, roadmap, action contract and Singapore case","Draft reduction-plan architecture"),
       ("17:45–18:30","0:45","Activities 7–9 and course synthesis",f"{sm['activities']['7']}–{sm['activities']['9']}","MR readiness, engagement design and capstone board review","Readiness plan, programme canvas and board-ready roadmap")])]
    for day,rows in schedules:
        d.add_page_break(); heading(d,f"Day {day}: {C.DAY_THEMES[day]}",1)
        table=d.add_table(rows=1,cols=6); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,h in enumerate(cols): table.rows[0].cells[i].text=h
        for row in rows: add_lp_row(table,row)
    d.add_page_break(); heading(d,"Assessment Schedule — Separate 2 Hours",1)
    table=d.add_table(rows=1,cols=5); table.style="Table Grid"
    for i,h in enumerate(["Duration","Instrument","Competency","Conditions","Evidence"]): table.rows[0].cells[i].text=h
    add_lp_row(table,["1:00","Written Assessment — 9 SAQs","K1–K9","Open book; individual", "Written responses and calculations"])
    add_lp_row(table,["1:00","Written Assessment — Case Study, 3 questions","A1–A3","Open book; individual", "Hotspot analysis, prioritisation and reduction plan"])
    heading(d,"Resources",1); bullets(d,["v7 course slide deck","v7 Learner Guide","Nine activity folders and datasets","SG Carbon Calculator","Source register and current primary-source checks"])
    path=os.path.join(OUT,f"LP-{C.SHORT_TITLE}-{C.VERSION}.docx"); d.save(path); return path

if __name__=="__main__":
    print("Saved",build_lg())
    print("Saved",build_lp())
